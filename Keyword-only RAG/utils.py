import os
from glob import glob
from tqdm import tqdm
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import re
import time

def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = text.replace('\\t', ' ')
    text = text.replace('\\n', ' ')
    text = ' '.join(text.split())
    return text

def process_multiple_documents(data_dir):
    documents = []
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", " ", ""]
    )
    
    # 收集所有符合模式的文件路径
    all_files = []
    for file_pattern in ["**/*.md", "**/*.docx", "**/*.pdf"]:
        all_files.extend(glob(os.path.join(data_dir, file_pattern), recursive=True))
    
    # 使用tqdm显示文件处理进度
    for file_path in tqdm(all_files, desc="处理文件", unit="个文件"):
        try:
            if file_path.endswith(".md"):
                with open(file_path, "r", encoding="utf-8") as file:
                    file_text = file.read().strip()
            elif file_path.endswith(".docx"):
                doc = DocxDocument(file_path)
                paragraphs = [para.text for para in doc.paragraphs]
                file_text = "\n".join(paragraphs).strip()
            elif file_path.endswith(".pdf"):
                pdf_reader = PdfReader(file_path)
                text_list = []
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_list.append(page.extract_text())
                file_text = "\n".join(text_list).strip()
            else:
                continue
            if file_text:
                file_text = clean_text(file_text)
                chunks = text_splitter.split_text(file_text)
                for chunk in chunks:
                    if chunk.strip():
                        documents.append(Document(
                            page_content=chunk,
                            metadata={"source": file_path}
                        ))
        except Exception as e:
            print(f"Error processing file {file_path}: {e}")
            continue
    print(f"Processed {len(documents)} document chunks from {data_dir}")
    return documents